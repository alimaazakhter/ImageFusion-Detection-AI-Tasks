from docx import Document
import re

# Mapping from template placeholders to LLM keys
TEMPLATE_TO_LLM_KEY = {
    'DATE_LOSS': "Date Taken",
    'INSURED_NAME': "Insured's Name",
    'INSURED_H_STREET': "Address of Loss",
    'CARRIER_NAME': "Carrier Name",
    'POLICY_NO': "Policy #",
    'SERVICE_PROVIDER': "Service Provider",
    'SERVICE_PROVIDER_ADDRESS': "Service Provider Address",
    'SERVICE_PROVIDER_PHONE': "Service Provider Phone",
    # Add more mappings as needed
    # 'TEMPLATE_KEY': 'LLM Key',
}

def fill_docx_template(template_file, key_value_pairs, output_path):
    """
    Fills a DOCX template with key-value pairs and saves the result to output_path.
    Supports placeholders in the form {{FieldName}} or [FieldName], case-insensitive.
    Uses TEMPLATE_TO_LLM_KEY to map template keys to LLM keys.
    Warns if a placeholder is not found in the extracted data.
    """
    try:
        doc = Document(template_file)
        # Build a case-insensitive key map
        key_map = {k.lower(): v for k, v in key_value_pairs.items()}
        # Patterns for both {{KEY}} and [KEY]
        patterns = [re.compile(r'\{\{(.*?)\}\}'), re.compile(r'\[(.*?)\]')]
        warnings = set()
        def replace_placeholders(text):
            for pattern in patterns:
                matches = pattern.findall(text)
                for match in matches:
                    key = match.strip()
                    # Try mapping first
                    mapped_key = TEMPLATE_TO_LLM_KEY.get(key.upper())
                    value = None
                    if mapped_key:
                        value = key_map.get(mapped_key.lower())
                    if value is None:
                        # Fallback to direct match
                        value = key_map.get(key.lower())
                    if value is not None:
                        text = text.replace(f"{{{{{key}}}}}", str(value))
                        text = text.replace(f"[{key}]", str(value))
                    else:
                        warnings.add(key)
            return text
        # Replace in paragraphs
        for para in doc.paragraphs:
            inline = para.runs
            for i in range(len(inline)):
                text = inline[i].text
                new_text = replace_placeholders(text)
                inline[i].text = new_text
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        inline = para.runs
                        for i in range(len(inline)):
                            text = inline[i].text
                            new_text = replace_placeholders(text)
                            inline[i].text = new_text
        doc.save(output_path)
        if warnings:
            print(f"Warning: The following placeholders were not found in the extracted data: {sorted(warnings)}")
        return True
    except Exception as e:
        print(f"Error filling DOCX template: {e}")
        return False 